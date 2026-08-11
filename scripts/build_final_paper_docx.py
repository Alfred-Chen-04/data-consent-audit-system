from __future__ import annotations

import argparse
from copy import deepcopy
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

REPLACEMENTS = {
    "Prepared: August 5, 2026": "Prepared: August 11, 2026",
    "Transparency and refusal effectiveness regressed": (
        "Refusal effectiveness regressed; transparency not assessed"
    ),
    "The later regression is documented; its cause is not.": (
        "The later regression is documented; its cause is not. Because the "
        "earlier record lacks a comparable transparency baseline, transparency "
        "direction is not assessed."
    ),
}

PROGRAM_LINE = "Program: CWRU 2026 Sponsored Summer Research Program"


def _replace_in_paragraph(paragraph: Paragraph) -> int:
    original = paragraph.text
    updated = original
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)
    if updated == original:
        return 0
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)
    return 1


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document(args.source)
    replacements = sum(_replace_in_paragraph(p) for p in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                replacements += sum(
                    _replace_in_paragraph(paragraph) for paragraph in cell.paragraphs
                )

    if replacements != 3:
        raise RuntimeError(f"Expected 3 replacements, made {replacements}")

    prepared = next(
        (paragraph for paragraph in document.paragraphs if paragraph.text.startswith("Prepared:")),
        None,
    )
    if prepared is None:
        raise RuntimeError("Prepared line was not found")
    program_xml = deepcopy(prepared._p)
    text_nodes = program_xml.xpath(".//w:t")
    if not text_nodes:
        raise RuntimeError("Prepared line has no text node")
    text_nodes[0].text = PROGRAM_LINE
    for node in text_nodes[1:]:
        node.text = ""
    prepared._p.addnext(program_xml)

    document.core_properties.title = "How Cookie Consent Interfaces Changed"
    document.core_properties.subject = "SSRP 2026 final paper"
    document.core_properties.modified = datetime.now(UTC)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)


if __name__ == "__main__":
    main()
